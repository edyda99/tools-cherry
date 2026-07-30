"""Hostile-input regression suite for docx_enhance passes.

Every case here either predates a pass (design guards) or reproduces a defect
found by the adversarial review of iteration 2 (link/image destruction, stream
misalignment, nested-ordered renumbering, OOXML sequence breaks). Run on every
loop iteration: venv/bin/python hostile_tests.py  — exit 0 only when all pass.
"""
import io
import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent.parent))

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt

import docx_enhance as de

FAILS = []


def check(cond, msg):
    if not cond:
        FAILS.append(msg)
        print("FAIL:", msg)


def run_pass(doc):
    buf = io.BytesIO()
    doc.save(buf)
    out = de.list_numbering(buf.getvalue())
    return Document(io.BytesIO(out)), out


def has_numpr(p):
    ppr = p._p.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:numPr")) is not None


def ilvl_of(p):
    el = p._p.find(qn("w:pPr")).find(qn("w:numPr")).find(qn("w:ilvl"))
    return int(el.get(qn("w:val")))


def numid_of(p):
    el = p._p.find(qn("w:pPr")).find(qn("w:numPr")).find(qn("w:numId"))
    return int(el.get(qn("w:val")))


def count_tag(docx_bytes, tag):
    import zipfile
    xml = zipfile.ZipFile(io.BytesIO(docx_bytes)).read("word/document.xml").decode()
    return xml.count("<" + tag)


def full_text(doc):
    """Every w:t in the body, including inside hyperlinks."""
    return "".join(t.text or "" for p in doc.paragraphs for t in p._p.iter(qn("w:t")))


# --- A: prose that must never convert, lists that must -----------------------
d = Document()
cases = ["3.14 is the ratio of circumference to diameter",
         "- a lone dash aside, prose not list",
         "2. starts at two so numbering must not touch it",
         "2026. was a fine year for the project",
         "1. one", "2. two", "3. three",
         "• bullet keeps its inline • dot",
         "10) parenthesised but starting at ten"]
for c in cases:
    d.add_paragraph(c)
r, _ = run_pass(d)
p = r.paragraphs
check(p[0].text == cases[0] and not has_numpr(p[0]), "A: 3.14 corrupted")
check(p[1].text == cases[1] and not has_numpr(p[1]), "A: lone dash converted")
check(p[2].text == cases[2] and not has_numpr(p[2]), "A: starts-at-two converted")
check(p[3].text == cases[3] and not has_numpr(p[3]), "A: year corrupted")
check(p[4].text == "one" and has_numpr(p[4]), "A: 1..3 not converted")
check(p[6].text == "three" and has_numpr(p[6]), "A: third item wrong")
check(p[7].text == "bullet keeps its inline • dot" and has_numpr(p[7]), "A: inline dot wrong")
check(p[8].text == cases[8] and not has_numpr(p[8]), "A: starts-at-ten converted")

d = Document()
for c in ["- first dash item", "- second dash item", "regular prose after"]:
    d.add_paragraph(c)
r, _ = run_pass(d)
check(r.paragraphs[0].text == "first dash item" and has_numpr(r.paragraphs[0]),
      "A: dash pair not converted")
check(not has_numpr(r.paragraphs[2]), "A: prose after converted")

# --- B: pdf2docx-shaped hyperlink (w:hyperlink nested INSIDE a w:r) ----------
d = Document()
para = d.add_paragraph()
para._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">• Start with the </w:t></w:r>'))
para._p.append(parse_xml(
    f'<w:r {nsdecls("w")}><w:rPr/><w:hyperlink {nsdecls("w")} w:anchor="x">'
    f'<w:r><w:t>tools-berry homepage</w:t></w:r></w:hyperlink></w:r>'))
para._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve"> for context.</w:t></w:r>'))
d.add_paragraph("• second item so the block is unambiguous")
r, out = run_pass(d)
check(count_tag(out, "w:hyperlink") == 1, "B: nested hyperlink deleted")
check("tools-berry homepage" in full_text(r), "B: link text deleted")
check(has_numpr(r.paragraphs[0]), "B: numPr missing")
check(full_text(r).startswith("Start with the "), "B: marker not stripped cleanly")

# --- C: standards-shaped hyperlink carrying the marker => skip entirely ------
d = Document()
para = d.add_paragraph()
para._p.append(parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="x">'
                         f'<w:r><w:t>1. Annual Report</w:t></w:r></w:hyperlink>'))
para._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve"> (PDF, 2 MB)</w:t></w:r>'))
para2 = d.add_paragraph()
para2._p.append(parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="y">'
                          f'<w:r><w:t>2. Board Minutes</w:t></w:r></w:hyperlink>'))
r, out = run_pass(d)
check(full_text(r) == "1. Annual Report (PDF, 2 MB)2. Board Minutes",
      "C: hyperlink-marker paragraph text corrupted: %r" % full_text(r))
check(not has_numpr(r.paragraphs[0]) and not has_numpr(r.paragraphs[1]),
      "C: hyperlink-marker paragraph got numPr")

# --- D: paragraph that is entirely one hyperlink => untouched ----------------
d = Document()
para = d.add_paragraph()
para._p.append(parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="x">'
                         f'<w:r><w:t>• Entirely linked item</w:t></w:r></w:hyperlink>'))
r, out = run_pass(d)
check(not has_numpr(r.paragraphs[0]) and "• Entirely linked item" in full_text(r),
      "D: whole-hyperlink paragraph modified")

# --- E/F/G: images and references survive ------------------------------------
d = Document()
para = d.add_paragraph("• chart below ")
para._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:drawing/></w:r>'))
d.add_paragraph("• second")
r, out = run_pass(d)
check(count_tag(out, "w:drawing") == 1, "E: drawing-only run deleted")
check(has_numpr(r.paragraphs[0]) and r.paragraphs[0].text == "chart below ",
      "E: marker strip wrong around image")

d = Document()
para = d.add_paragraph()
para._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">• </w:t><w:drawing/></w:r>'))
para._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">caption text</w:t></w:r>'))
d.add_paragraph("• second")
r, out = run_pass(d)
check(count_tag(out, "w:drawing") == 1, "F: drawing sharing marker run deleted")
check("caption text" in full_text(r), "F: caption lost")

d = Document()
para = d.add_paragraph("• see note")
para._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:footnoteReference w:id="2"/></w:r>'))
d.add_paragraph("• second")
r, out = run_pass(d)
check(count_tag(out, "w:footnoteReference") == 1, "G: footnoteReference run deleted")

# --- H: ordered nested under a bullet must still render decimal --------------
d = Document()
d.add_paragraph("• Before you begin")
for t in ["1. Install the vendor driver", "2. Reboot the machine", "3. Run the self-test"]:
    pp = d.add_paragraph(t)
    pp.paragraph_format.left_indent = Pt(42)
r, out = run_pass(d)
steps = r.paragraphs[1:4]
check(all(has_numpr(s) for s in steps), "H: nested ordered not converted")
lvls = {ilvl_of(s) for s in steps}
check(len(lvls) == 1, "H: ordered stretch split across levels")
import zipfile
numxml = zipfile.ZipFile(io.BytesIO(out)).read("word/numbering.xml").decode()
check('w:val="lowerLetter"' not in numxml and 'w:val="lowerRoman"' not in numxml,
      "H: ordered levels not all decimal")

# --- I: indent jitter inside a flat ordered list => one level ----------------
d = Document()
for i, t in enumerate(["1. first", "2. second", "3. third"]):
    pp = d.add_paragraph(t)
    pp.paragraph_format.left_indent = Pt(20 if i == 1 else 0)
r, out = run_pass(d)
check(len({ilvl_of(pp) for pp in r.paragraphs}) == 1, "I: jitter split the sequence")

# --- J: numPr must land after keepNext/pageBreakBefore/widowControl ----------
d = Document()
pp = d.add_paragraph("• formatted item")
pp.paragraph_format.keep_with_next = True
pp.paragraph_format.page_break_before = True
pp.paragraph_format.widow_control = True
d.add_paragraph("• second")
r, out = run_pass(d)
ppr = r.paragraphs[0]._p.find(qn("w:pPr"))
tags = [c.tag.split("}")[1] for c in ppr]
check("numPr" in tags and tags.index("numPr") > tags.index("keepNext"),
      "J: numPr precedes keepNext in pPr sequence: %s" % tags)

# --- K: w:num insertion respects numIdMacAtCleanup ---------------------------
d = Document()
d.add_paragraph("• one")
d.add_paragraph("• two")
numbering = de._numbering_root(d)
numbering.append(parse_xml(f'<w:numIdMacAtCleanup {nsdecls("w")} w:val="9"/>'))
r, out = run_pass(d)
numxml = zipfile.ZipFile(io.BytesIO(out)).read("word/numbering.xml").decode()
check(numxml.rstrip().endswith("numIdMacAtCleanup w:val=\"9\"/></w:numbering>")
      or numxml.find("<w:num ", numxml.find("numIdMacAtCleanup")) == -1,
      "K: w:num appended after numIdMacAtCleanup")

# --- L: w:ind @w:start honoured for levels -----------------------------------
d = Document()
for t, tw in [("• top", 0), ("◦ nested", 720)]:
    pp = d.add_paragraph(t)
    ppr = pp._p.get_or_add_pPr()
    ppr.append(parse_xml(f'<w:ind {nsdecls("w")} w:start="{tw}"/>'))
r, out = run_pass(d)
check(ilvl_of(r.paragraphs[0]) == 0 and ilvl_of(r.paragraphs[1]) == 1,
      "L: @w:start indents not levelled")

# --- M: noBreakHyphen survives a strip in the same run -----------------------
d = Document()
para = d.add_paragraph()
para._p.append(parse_xml(
    f'<w:r {nsdecls("w")}><w:t>1.</w:t><w:tab/><w:t>state</w:t>'
    f'<w:noBreakHyphen/><w:t>of-the-art</w:t></w:r>'))
d.add_paragraph("2. two")
r, out = run_pass(d)
check(count_tag(out, "w:noBreakHyphen") == 1, "M: noBreakHyphen destroyed")
check(r.paragraphs[0].text == "state-of-the-art", "M: text wrong: %r" % r.paragraphs[0].text)

# --- N: marker-only paragraph stays untouched --------------------------------
d = Document()
para = d.add_paragraph()
para._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">• </w:t></w:r>'))
d.add_paragraph("• real item")
d.add_paragraph("• real item two")
r, out = run_pass(d)
check(not has_numpr(r.paragraphs[0]) and r.paragraphs[0].text == "• ",
      "N: marker-only paragraph converted")

# --- O: regression set the reviewers verified --------------------------------
d = Document()
tbl = d.add_table(rows=1, cols=2)
tbl.rows[0].cells[0].text = "• cell bullet"
tbl.rows[0].cells[1].text = "1. cell ordered"
before = io.BytesIO()
d.save(before)
out = de.list_numbering(before.getvalue())
check(out == before.getvalue(), "O: table-cell lists modified")

d = Document()
r, out = run_pass(d)  # empty document
check(len(r.paragraphs) <= 1, "O: empty doc changed shape")

d = Document()
for t in ["1. a", "2. b", "3. c"]:
    d.add_paragraph(t)
d.add_paragraph("prose between")
for t in ["1. x", "2. y", "3. z"]:
    d.add_paragraph(t)
r, out = run_pass(d)
check(numid_of(r.paragraphs[0]) != numid_of(r.paragraphs[4]),
      "O: two ordered lists share a numId (no restart)")

d = Document()
for i in range(1, 13):
    d.add_paragraph(f"{i}. item {i}")
r, out = run_pass(d)
check(all(has_numpr(pp) for pp in r.paragraphs), "O: 1..12 not fully converted")

d = Document()
for t, ind in [("• l0", 0), ("◦ l1", 36), ("▪ l2", 72), ("▪ l3", 108)]:
    pp = d.add_paragraph(t)
    pp.paragraph_format.left_indent = Pt(ind)
r, out = run_pass(d)
check([ilvl_of(pp) for pp in r.paragraphs] == [0, 1, 2, 2], "O: level cap wrong")

# --- P: hyperlink first, marker in a later direct run => skip ----------------
d = Document()
para = d.add_paragraph()
para._p.append(parse_xml(f'<w:hyperlink {nsdecls("w")} w:anchor="x">'
                         f'<w:r><w:t>See </w:t></w:r></w:hyperlink>'))
para._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t>2. something</w:t></w:r>'))
r, out = run_pass(d)
check(not has_numpr(r.paragraphs[0]) and full_text(r) == "See 2. something",
      "P: mid-paragraph marker after hyperlink converted")

# === header/footer pass ======================================================
import fitz


def make_pdf(pages):
    """pages: list of [(y, text) or (y, text, x), ...] on A4."""
    pdf = fitz.open()
    for lines in pages:
        pg = pdf.new_page(width=595, height=842)
        for entry in lines:
            y, t = entry[0], entry[1]
            x = entry[2] if len(entry) > 2 else 72
            pg.insert_text((x, y), t, fontsize=9)
    return pdf


def run_hf(doc, pdf):
    buf = io.BytesIO()
    doc.save(buf)
    out = de.header_footer_parts(buf.getvalue(), pdf)
    return Document(io.BytesIO(out)), out


def hf_part_text(docx_bytes, which):
    z = zipfile.ZipFile(io.BytesIO(docx_bytes))
    texts = []
    for n in z.namelist():
        if n.startswith(f"word/{which}"):
            import xml.etree.ElementTree as ET
            root = ET.fromstring(z.read(n))
            texts.extend(t.text or "" for t in root.iter(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"))
    return "".join(texts)


# HF1: repeated band furniture moves into real parts
pdf = make_pdf([[(30, "Acme Internal"), (400, "content one"), (820, "Confidential")],
                [(30, "Acme Internal"), (400, "content two"), (820, "Confidential")]])
d = Document()
for t in ["Acme Internal", "content one", "Confidential",
          "Acme Internal", "content two", "Confidential"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
body = [p.text for p in r.paragraphs if p.text.strip()]
check(body == ["content one", "content two"], "HF1: body not cleaned: %s" % body)
check("Acme Internal" in hf_part_text(out, "header"), "HF1: header part missing text")
check("Confidential" in hf_part_text(out, "footer"), "HF1: footer part missing text")

# HF2: varying page numbers stay (no exact repeat)
pdf = make_pdf([[(400, "content one"), (820, "1")], [(400, "content two"), (820, "2")]])
d = Document()
for t in ["content one", "1", "content two", "2"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
check([p.text for p in r.paragraphs] == ["content one", "1", "content two", "2"],
      "HF2: varying page numbers touched")

# HF3: repeated text mid-page is content, not furniture
pdf = make_pdf([[(400, "Same refrain"), (500, "other a")], [(400, "Same refrain"), (500, "other b")]])
d = Document()
for t in ["Same refrain", "other a", "Same refrain", "other b"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if p.text == "Same refrain") == 2,
      "HF3: mid-page repeat treated as furniture")

# HF4: single page never converts
pdf = make_pdf([[(30, "Lone header"), (400, "content")]])
d = Document()
for t in ["Lone header", "content"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
check([p.text for p in r.paragraphs] == ["Lone header", "content"], "HF4: 1-page doc touched")

# HF5: more body matches than pages => ambiguous, skip
pdf = make_pdf([[(30, "Motto"), (400, "a")], [(30, "Motto"), (400, "b")]])
d = Document()
for t in ["Motto", "a", "Motto", "b", "Motto", "Motto"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if p.text == "Motto") == 4, "HF5: over-matched text removed")

# HF6: a match carrying a section break is never removed
pdf = make_pdf([[(30, "Banner"), (400, "a")], [(30, "Banner"), (400, "b")]])
d = Document()
d.add_paragraph("Banner")
d.add_paragraph("a")
p2 = d.add_paragraph("Banner")
p2._p.get_or_add_pPr().append(parse_xml(f"<w:sectPr {nsdecls('w')}/>"))
d.add_paragraph("b")
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if p.text == "Banner") == 2, "HF6: sectPr paragraph removed")

# HF7: furniture must cover EVERY page — 2 of 3 declines, 3 of 3 converts
pdf = make_pdf([[(400, "x"), (820, "Draft")], [(400, "y"), (820, "Draft")], [(400, "z")]])
d = Document()
for t in ["x", "Draft", "y", "Draft", "z"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if p.text == "Draft") == 2 and
      "Draft" not in hf_part_text(out, "footer"), "HF7a: partial-coverage footer converted")
pdf = make_pdf([[(400, "x"), (820, "Draft")], [(400, "y"), (820, "Draft")], [(400, "z"), (820, "Draft")]])
d = Document()
for t in ["x", "Draft", "y", "Draft", "z", "Draft"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
check("Draft" in hf_part_text(out, "footer") and
      all(p.text != "Draft" for p in r.paragraphs), "HF7b: full-coverage footer not converted")

# HF8: cover-title collision — header on pages 2-6 only, title on page 1 => decline
pdf = make_pdf([[(300, "Quarterly Report")]] +
               [[(30, "Quarterly Report"), (400, f"body {i}")] for i in range(5)])
d = Document()
d.add_paragraph("Quarterly Report")
for i in range(5):
    d.add_paragraph("Quarterly Report")
    d.add_paragraph(f"body {i}")
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if p.text == "Quarterly Report") == 6,
      "HF8: cover title deleted with running headers")

# HF9: under-match (one occurrence merged/missing) => decline, no half-removal
pdf = make_pdf([[(30, "ACME Holdings"), (400, "a")], [(30, "ACME Holdings"), (400, "b")],
                [(30, "ACME Holdings"), (400, "c")]])
d = Document()
for t in ["ACME Holdings", "a", "ACME Holdings", "b", "ACME Holdings and page text", "c"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if "ACME Holdings" in p.text) == 3,
      "HF9: half-removal on merged occurrence")
check("ACME" not in hf_part_text(out, "header"), "HF9: header installed despite mismatch")

# HF10: text also inside a table cell => decline
pdf = make_pdf([[(30, "Classified"), (400, "a")], [(30, "Classified"), (400, "b")]])
d = Document()
d.add_paragraph("Classified")
d.add_paragraph("a")
tbl = d.add_table(rows=1, cols=1)
tbl.rows[0].cells[0].text = "Classified"
d.add_paragraph("Classified")
d.add_paragraph("b")
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if p.text == "Classified") == 2,
      "HF10: cell-shadowed furniture removed")

# HF11: same text qualifying in both bands => decline
pdf = make_pdf([[(30, "Everywhere"), (820, "Everywhere"), (400, "a")],
                [(30, "Everywhere"), (820, "Everywhere"), (400, "b")]])
d = Document()
for t in ["Everywhere", "a", "Everywhere", "Everywhere", "b", "Everywhere"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if p.text == "Everywhere") == 4,
      "HF11: both-band text swept")

# HF12: bookmark/footnote-bearing furniture => decline (range pairs stay whole)
pdf = make_pdf([[(30, "Marked"), (400, "a")], [(30, "Marked"), (400, "b")]])
d = Document()
pm = d.add_paragraph("Marked")
pm._p.insert(0, parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="7" w:name="top"/>'))
d.add_paragraph("a")
d.add_paragraph("Marked")
d.add_paragraph("b")
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if p.text == "Marked") == 2, "HF12: bookmark furniture moved")

pdf = make_pdf([[(30, "Noted"), (400, "a")], [(30, "Noted"), (400, "b")]])
d = Document()
pn = d.add_paragraph("Noted")
pn._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:footnoteReference w:id="2"/></w:r>'))
d.add_paragraph("a")
d.add_paragraph("Noted")
d.add_paragraph("b")
r, out = run_hf(d, pdf)
check("Noted" not in hf_part_text(out, "header"), "HF12b: footnote ref moved into header")

# HF13: existing header part => decline entirely; and the pass is idempotent
pdf = make_pdf([[(30, "Fresh"), (400, "a")], [(30, "Fresh"), (400, "b")]])
d = Document()
sec = d.sections[0]
sec.header.is_linked_to_previous = False
sec.header.paragraphs[0].text = "COMPANY LETTERHEAD - keep me"
for t in ["Fresh", "a", "Fresh", "b"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if p.text == "Fresh") == 2 and
      "keep me" in hf_part_text(out, "header"), "HF13: existing header not respected")

d = Document()
for t in ["Fresh", "a", "Fresh", "b"]:
    d.add_paragraph(t)
buf = io.BytesIO()
d.save(buf)
once = de.header_footer_parts(buf.getvalue(), pdf)
twice = de.header_footer_parts(once, pdf)
check(once == twice, "HF13b: pass not idempotent")

# HF14: titlePg set => decline (default header would skip page 1)
pdf = make_pdf([[(30, "Banner"), (400, "a")], [(30, "Banner"), (400, "b")]])
d = Document()
d.sections[0]._sectPr.append(parse_xml(f'<w:titlePg {nsdecls("w")}/>'))
for t in ["Banner", "a", "Banner", "b"]:
    d.add_paragraph(t)
r, out = run_hf(d, pdf)
check(sum(1 for p in r.paragraphs if p.text == "Banner") == 2, "HF14: titlePg doc converted")

# HF15: restarting ordered lists made adjacent still convert, separately
d = Document()
for t in ["1. alpha", "2. beta", "3. gamma", "1. delta", "2. epsilon", "3. zeta"]:
    d.add_paragraph(t)
r, out = run_pass(d)
check(all(has_numpr(pp) for pp in r.paragraphs), "HF15: adjacent restarting lists dropped")
check(numid_of(r.paragraphs[0]) != numid_of(r.paragraphs[3]),
      "HF15: restart shares a numId (would renumber 4,5,6)")

# HF16: small outline docs keep their headings (share-bail exemption)
d = Document()
for t in ["First Heading", "Second Heading", "Third Heading",
          "a single body paragraph long enough that the body size vote lands on "
          "eleven points, the way a real short outline document reads"]:
    pp = d.add_paragraph()
    run = pp.add_run(t)
    run.bold = "Heading" in t
    run.font.size = Pt(16 if "Heading" in t else 11)
buf = io.BytesIO()
d.save(buf)
res = Document(io.BytesIO(de.heading_styles(buf.getvalue())))
styled = [p.text for p in res.paragraphs if (p.style.name or "").startswith("Heading")]
check(len(styled) == 3, "HF16: small-doc headings bailed: %s" % styled)

# === paragraph reflow pass ===================================================

def run_reflow(doc, pdf):
    buf = io.BytesIO()
    doc.save(buf)
    out = de.paragraph_reflow(buf.getvalue(), pdf)
    return Document(io.BytesIO(out)), out


# R1: fragments of one PDF block merge (lines 12pt apart share a block)
pdf = make_pdf([[(100, "the quick brown fox jumps across the sleeping meadow and"),
                 (112, "over the lazy dog.")]])
d = Document()
d.add_paragraph("the quick brown fox jumps across the sleeping meadow and")
d.add_paragraph("over the lazy dog.")
r, out = run_reflow(d, pdf)
texts = [p.text for p in r.paragraphs if p.text.strip()]
check(texts == ["the quick brown fox jumps across the sleeping meadow and over the lazy dog."],
      "R1: same-block fragments not merged: %s" % texts)

# R2: separate PDF blocks stay separate paragraphs
pdf = make_pdf([[(100, "the quick brown fox jumps"), (400, "over the lazy dog.")]])
d = Document()
d.add_paragraph("the quick brown fox jumps")
d.add_paragraph("over the lazy dog.")
r, out = run_reflow(d, pdf)
check(len([p for p in r.paragraphs if p.text.strip()]) == 2, "R2: separate blocks merged")

# R3: styled / numbered / section-break paragraphs never merge
pdf = make_pdf([[(100, "alpha beta gamma"), (112, "delta epsilon.")]])
d = Document()
d.add_paragraph("alpha beta gamma")
pb = d.add_paragraph("delta epsilon.")
pb._p.get_or_add_pPr().append(parse_xml(f"<w:sectPr {nsdecls('w')}/>"))
r, out = run_reflow(d, pdf)
check(len([p for p in r.paragraphs if p.text.strip()]) == 2, "R3: sectPr paragraph merged")

# R4: typographic (U+2010) line-break hyphens heal; ASCII ones never do
pdf = fitz.open()
pg = pdf.new_page(width=595, height=842)
pg.insert_font(fontname="hv", fontfile="/System/Library/Fonts/Helvetica.ttc")
pg.insert_text((72, 100), "the grand equip‐", fontsize=9, fontname="hv")
pg.insert_text((72, 112), "ment failed early on the very first day of trials.",
               fontsize=9, fontname="hv")
d = Document()
d.add_paragraph("the grand equip‐ment failed early on the very first day of trials.")
r, out = run_reflow(d, pdf)
check(r.paragraphs[0].text == "the grand equipment failed early on the very first day of trials.",
      "R4: typographic hyphen not healed: %r" % r.paragraphs[0].text)
pdf = make_pdf([[(100, "the grand equip-"), (112, "ment failed early."),
                 (400, "equipment budgets grew.")]])
d = Document()
d.add_paragraph("the grand equip-ment failed early.")
d.add_paragraph("equipment budgets grew.")
r, out = run_reflow(d, pdf)
check("equip-ment" in r.paragraphs[0].text, "R4b: ASCII hyphen fused")

# R11: a compound word broken on its own hyphen at a line end stays hyphenated
pdf = make_pdf([[(100, "that fact is well-"), (112, "known to every reader.")]])
d = Document()
d.add_paragraph("that fact is well-known to every reader.")
r, out = run_reflow(d, pdf)
check(r.paragraphs[0].text == "that fact is well-known to every reader.",
      "R11: compound hyphen fused: %r" % r.paragraphs[0].text)

# R5: a real hyphen with no line-break evidence stays
pdf = make_pdf([[(100, "a well- known fact stands.")]])
d = Document()
d.add_paragraph("a well- known fact stands.")
r, out = run_reflow(d, pdf)
check(r.paragraphs[0].text == "a well- known fact stands.", "R5: legit hyphen removed")

# R6: a table between paragraphs breaks adjacency
pdf = make_pdf([[(100, "alpha beta gamma"), (112, "delta epsilon.")]])
d = Document()
d.add_paragraph("alpha beta gamma")
d.add_table(rows=1, cols=1).rows[0].cells[0].text = "cell"
d.add_paragraph("delta epsilon.")
r, out = run_reflow(d, pdf)
check(len([p for p in r.paragraphs if p.text.strip()]) == 2, "R6: merged across a table")

# R7: idempotent on bytes
pdf = make_pdf([[(100, "the quick brown fox jumps"), (112, "over the lazy dog.")]])
d = Document()
d.add_paragraph("the quick brown fox jumps")
d.add_paragraph("over the lazy dog.")
buf = io.BytesIO()
d.save(buf)
once = de.paragraph_reflow(buf.getvalue(), pdf)
twice = de.paragraph_reflow(once, pdf)
check(once == twice, "R7: reflow not idempotent")

# R8: no joining across page breaks (glued page-boundary text corrupted docs)
pdf = make_pdf([[(800, "the meeting ran long and")], [(50, "nobody minded at all.")]])
d = Document()
d.add_paragraph("the meeting ran long and")
d.add_paragraph("nobody minded at all.")
r, out = run_reflow(d, pdf)
check(len([p.text for p in r.paragraphs if p.text.strip()]) == 2,
      "R8: merged across a page break")

# R21: a compound broken once at a line end never fuses, even with the fused
# word used elsewhere ("re-form" vs "reform" inverts meaning)
pdf = make_pdf([[(100, "the panel would have to re-"),
                 (112, "form before the review could restart in earnest."),
                 (300, "the reform of the charter is complete.")]])
d = Document()
d.add_paragraph("the panel would have to re-form before the review could restart in earnest.")
d.add_paragraph("the reform of the charter is complete.")
r, out = run_reflow(d, pdf)
check("re-form" in r.paragraphs[0].text, "R21: compound fused into its homograph")

# R22: lowercase unterminated paragraphs in separate blocks never weld
pdf = make_pdf([[(100, "max_retries the number of attempts made before the call fails and"),
                 (112, "an error is returned to the caller"),
                 (127, "timeout the number of seconds the client waits before closing"),
                 (139, "the socket entirely"),
                 (154, "verify_tls whether the certificate chain is validated at all")]])
d = Document()
d.add_paragraph("max_retries the number of attempts made before the call fails and an error is returned to the caller")
d.add_paragraph("timeout the number of seconds the client waits before closing the socket entirely")
d.add_paragraph("verify_tls whether the certificate chain is validated at all")
r, out = run_reflow(d, pdf)
check(len([p for p in r.paragraphs if p.text.strip()]) == 3,
      "R22: config paragraphs welded: %d" % len([p for p in r.paragraphs if p.text.strip()]))

# R9: a paragraph that already equals a full PDF paragraph absorbs nothing
pdf = make_pdf([[(100, "first sentence stands alone"), (400, "second one also complete.")]])
d = Document()
d.add_paragraph("first sentence stands alone")
d.add_paragraph("second one also complete.")
r, out = run_reflow(d, pdf)
check(len([p for p in r.paragraphs if p.text.strip()]) == 2, "R9: complete para absorbed next")

# R10: empty spacer between fragments rides along; a sectPr spacer blocks
pdf = make_pdf([[(100, "columns split this long sentence apart without any real warning"),
                 (112, "into two ragged pieces.")]])
d = Document()
d.add_paragraph("columns split this long sentence apart without any real warning")
d.add_paragraph("")
d.add_paragraph("into two ragged pieces.")
r, out = run_reflow(d, pdf)
check([p.text for p in r.paragraphs if p.text.strip()] ==
      ["columns split this long sentence apart without any real warning into two ragged pieces."],
      "R10: spacer merge missed")
d = Document()
d.add_paragraph("columns split this long sentence apart without any real warning")
ps = d.add_paragraph("")
ps._p.get_or_add_pPr().append(parse_xml(f"<w:sectPr {nsdecls('w')}/>"))
d.add_paragraph("into two ragged pieces.")
r, out = run_reflow(d, pdf)
check(len([p for p in r.paragraphs if p.text.strip()]) == 2, "R10b: merged across sectPr spacer")

# R12: two sentences single-spaced in one block stay two paragraphs
pdf = make_pdf([[(100, "The first policy took effect in March."),
                 (112, "Employees must file the new form.")]])
d = Document()
d.add_paragraph("The first policy took effect in March.")
d.add_paragraph("Employees must file the new form.")
r, out = run_reflow(d, pdf)
check(len([p for p in r.paragraphs if p.text.strip()]) == 2, "R12: sentence pair merged")

# R13: a lowercase word list is not a wrapped paragraph
pdf = make_pdf([[(100, "apples"), (114, "bananas"), (128, "cherries"), (142, "dates")]])
d = Document()
for t in ["apples", "bananas", "cherries", "dates"]:
    d.add_paragraph(t)
r, out = run_reflow(d, pdf)
check(len([p for p in r.paragraphs if p.text.strip()]) == 4, "R13: word list merged")

# R14: a non-Latin paragraph between Latin ones survives and blocks merging
pdf = make_pdf([[(100, "the reading room stays open"), (400, "until the last train leaves.")]])
d = Document()
d.add_paragraph("the reading room stays open")
d.add_paragraph("читальный зал открыт")
d.add_paragraph("until the last train leaves.")
r, out = run_reflow(d, pdf)
texts = [p.text for p in r.paragraphs if p.text.strip()]
check(len(texts) == 3 and "читальный зал открыт" in texts, "R14: non-Latin paragraph lost: %s" % texts)

# R15/R16: image-only and page-break paragraphs are content, never spacers
pdf = make_pdf([[(100, "the report continues past the small inline"),
                 (112, "figure and finishes on this line.")]])
d = Document()
d.add_paragraph("the report continues past the small inline")
pimg = d.add_paragraph()
pimg._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:drawing/></w:r>'))
d.add_paragraph("figure and finishes on this line.")
r, out = run_reflow(d, pdf)
check(count_tag(out, "w:drawing") == 1, "R15: image paragraph swept")
check(len([p for p in r.paragraphs]) == 3, "R15b: merged across image paragraph")
d = Document()
d.add_paragraph("the report continues past the small inline")
pbr = d.add_paragraph()
pbr._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:br w:type="page"/></w:r>'))
d.add_paragraph("figure and finishes on this line.")
r, out = run_reflow(d, pdf)
check('w:type="page"' in zipfile.ZipFile(io.BytesIO(out)).read("word/document.xml").decode(),
      "R16: page break swept")

# R17: a form the document also hyphenates mid-line is spelling, never fused
pdf = make_pdf([[(100, "contractors are asked to re-"),
                 (112, "sign the schedule every season and the"),
                 (124, "manager may also re-sign the cover page.")]])
d = Document()
d.add_paragraph("contractors are asked to re-sign the schedule every season and the "
                "manager may also re-sign the cover page.")
r, out = run_reflow(d, pdf)
check(r.paragraphs[0].text.count("re-sign") == 2,
      "R17: interior-evidenced hyphen fused: %r" % r.paragraphs[0].text)

# R18/R19: joiner space carries no underline and never doubles whitespace
pdf = make_pdf([[(100, "the final section of the annual report was written and"),
                 (112, "was revised again last week.")]])
d = Document()
pa = d.add_paragraph()
run = pa.add_run("the final section of the annual report was written and")
run.underline = True
d.add_paragraph("was revised again last week.")
r, out = run_reflow(d, pdf)
merged = r.paragraphs[0]
check(merged.text == "the final section of the annual report was written and was revised again last week.",
      "R18: merge text wrong: %r" % merged.text)
check(not any(run.text == " " and run.underline for run in merged.runs),
      "R18b: joiner space underlined")
d = Document()
d.add_paragraph("the final section of the annual report was written and ")
d.add_paragraph("was revised again last week.")
r, out = run_reflow(d, pdf)
check("and  was" not in r.paragraphs[0].text, "R19: doubled joiner space")

# R20: a first-line indent inside one block separates paragraphs
pdf = make_pdf([[(100, "the committee met for hours on the budget and"),
                 (112, "the debate ran long without a formal close"),
                 (124, "next quarter brought entirely new rules", 90)]])
d = Document()
d.add_paragraph("the committee met for hours on the budget and the debate ran long without a formal close")
d.add_paragraph("next quarter brought entirely new rules")
r, out = run_reflow(d, pdf)
check(len([p for p in r.paragraphs if p.text.strip()]) == 2, "R20: merged across indent")

# === span-boundary space repair ==============================================

def run_space(doc, pdf):
    buf = io.BytesIO()
    doc.save(buf)
    out = de.span_space_repair(buf.getvalue(), pdf)
    return Document(io.BytesIO(out)), out


# S1: evidenced lost space at a run boundary is restored
pdf = make_pdf([[(100, "Start with the deployment checklist before touching anything.")]])
d = Document()
p = d.add_paragraph()
p.add_run("Start with the")
p.add_run("deployment checklist before touching anything.")
r, out = run_space(d, pdf)
check(r.paragraphs[0].text == "Start with the deployment checklist before touching anything.",
      "S1: lost space not restored: %r" % r.paragraphs[0].text)

# S2: a word split across styled runs is never broken apart
pdf = make_pdf([[(100, "an unbelievable outcome was reported by everyone involved.")]])
d = Document()
p = d.add_paragraph()
p.add_run("an un")
p.add_run("believable outcome was reported by everyone involved.")
r, out = run_space(d, pdf)
check("un believable" not in r.paragraphs[0].text, "S2: intra-word split spaced")

# S3: fused form that is also a real PDF word stays untouched
pdf = make_pdf([[(100, "a round of talks began, and around the corner more waited.")]])
d = Document()
p = d.add_paragraph()
p.add_run("a")
p.add_run("round of talks began, and around the corner more waited.")
r, out = run_space(d, pdf)
check(r.paragraphs[0].text.startswith("around of talks"), "S3: ambiguous fusion modified")

# S4: already-spaced seams and seams across breaks stay untouched
pdf = make_pdf([[(100, "plain text with the deployment checklist ready.")]])
d = Document()
p = d.add_paragraph()
p.add_run("plain text with the ")
p.add_run("deployment checklist ready.")
buf = io.BytesIO()
d.save(buf)
check(de.span_space_repair(buf.getvalue(), pdf) == buf.getvalue(), "S4: spaced seam modified")

# S5: the pdf2docx hyperlink-wrapper shape is repaired through the nesting
pdf = make_pdf([[(100, "Start with the deployment checklist before touching production.")]])
d = Document()
p = d.add_paragraph()
p.add_run("Start with the")
p._p.append(parse_xml(
    f'<w:r {nsdecls("w")}><w:rPr/><w:hyperlink {nsdecls("w")} w:anchor="x">'
    f'<w:r><w:t>deployment checklist</w:t></w:r></w:hyperlink></w:r>'))
p._p.append(parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve"> before touching production.</w:t></w:r>'))
r, out = run_space(d, pdf)
check(count_tag(out, "w:hyperlink") == 1 and
      full_text(r) == "Start with the deployment checklist before touching production.",
      "S5: hyperlink seam not repaired: %r" % full_text(r))

# S6: idempotent
r2out = de.span_space_repair(out, pdf)
check(r2out == out, "S6: span repair not idempotent")

# S7: punctuation at the seam always declines, even with a tempting bigram
pdf = make_pdf([[(100, "Escalations go to the on-call engineer first thing."),
                 (114, "Every engineer takes a turn on call each quarter.")]])
d = Document()
p = d.add_paragraph()
p.add_run("Escalations go to the on")
p.add_run("-call engineer first thing.")
r, out = run_space(d, pdf)
check("on -call" not in r.paragraphs[0].text and "on-call" in r.paragraphs[0].text,
      "S7: hyphen seam spaced: %r" % r.paragraphs[0].text)

# S8: standards ids, ranges and apostrophe names decline
pdf = make_pdf([[(100, "Certified to ISO-9001 since 2018 and the ISO 9001 audit is annual."),
                 (114, "Coverage 2019-2024 inclusive; years 2019 2024 compared."),
                 (128, "O'Brien and O Brien both signed the register.")]])
d = Document()
p1 = d.add_paragraph()
p1.add_run("Certified to ISO")
p1.add_run("-9001 since 2018.")
p2 = d.add_paragraph()
p2.add_run("Coverage 2019")
p2.add_run("-2024 inclusive.")
p3 = d.add_paragraph()
p3.add_run("O'")
p3.add_run("Brien signed.")
r, out = run_space(d, pdf)
check("ISO-9001" in r.paragraphs[0].text, "S8a: ISO id spaced")
check("2019-2024" in r.paragraphs[1].text, "S8b: year range spaced")
check("O'Brien" in r.paragraphs[2].text, "S8c: apostrophe name spaced")

# S9: CJK line-wrap seams never receive a space
pdf = fitz.open()
pg = pdf.new_page(width=595, height=842)
pg.insert_font(fontname="cjk", fontfile="/System/Library/Fonts/Hiragino Sans GB.ttc")
pg.insert_text((72, 100), "日本語のテキス", fontsize=9, fontname="cjk")
pg.insert_text((72, 114), "トです", fontsize=9, fontname="cjk")
d = Document()
p = d.add_paragraph()
p.add_run("日本語のテキス")
p.add_run("トです")
buf = io.BytesIO()
d.save(buf)
check(de.span_space_repair(buf.getvalue(), pdf) == buf.getvalue(), "S9: CJK seam spaced")

# S11: a seam word with an interior hyphen is never probed as its substring
pdf = make_pdf([[(100, "The X-Rayscanner model 7 shipped this week to the lab."),
                 (114, "Each ray scanner was recalibrated on site by the crew.")]])
d = Document()
p = d.add_paragraph()
p.add_run("The X-Ray")
p.add_run("scanner model 7 shipped this week to the lab.")
r, out = run_space(d, pdf)
check("X-Rayscanner" in r.paragraphs[0].text,
      "S11: hyphen-interior token split: %r" % r.paragraphs[0].text)

# S12: accented seam words are whole tokens too, never their ASCII substrings
pdf = fitz.open()
pg = pdf.new_page(width=595, height=842)
pg.insert_font(fontname="hv", fontfile="/System/Library/Fonts/Helvetica.ttc")
pg.insert_text((72, 100), "Zürichbank AG posted quarterly results this morning.",
               fontsize=9, fontname="hv")
pg.insert_text((72, 114), "It is a rich bank with deep reserves and long history.",
               fontsize=9, fontname="hv")
d = Document()
p = d.add_paragraph()
p.add_run("Zürich")
p.add_run("bank AG posted quarterly results this morning.")
r, out = run_space(d, pdf)
check("Zürichbank" in r.paragraphs[0].text,
      "S12: accented token split: %r" % r.paragraphs[0].text)

# S10: drawing/text-box content is out of scope
pdf = make_pdf([[(100, "Charton call duty roster for the on call rotation.")]])
d = Document()
p = d.add_paragraph("Body text mentioning Charton")
p._p.append(parse_xml(
    f'<w:r {nsdecls("w")}><w:drawing><wps:txbx xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
    f'<w:txbxContent><w:p><w:r><w:t>on</w:t></w:r><w:r><w:t>call duty</w:t></w:r></w:p>'
    f'</w:txbxContent></wps:txbx></w:drawing></w:r>'))
buf = io.BytesIO()
d.save(buf)
out2 = de.span_space_repair(buf.getvalue(), pdf)
import xml.etree.ElementTree as _ET
doc_xml = _ET.fromstring(zipfile.ZipFile(io.BytesIO(out2)).read("word/document.xml"))
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
box = doc_xml.find(f".//{W_NS}txbxContent")
box_ts = [t.text for t in box.iter(f"{W_NS}t")]
check(box_ts == ["on", "call duty"], "S10: text-box seam edited: %s" % box_ts)

# ---- hyperlink_unnest cases (L) --------------------------------------------

W_MAIN = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def run_unnest(doc):
    buf = io.BytesIO()
    doc.save(buf)
    out = de.hyperlink_unnest(buf.getvalue())
    return Document(io.BytesIO(out)), out


def _hyperlinks(out_bytes):
    root = _ET.fromstring(zipfile.ZipFile(io.BytesIO(out_bytes)).read("word/document.xml"))
    return root, root.findall(f".//{W_MAIN}hyperlink")


# L1: nested hyperlink lifts to paragraph level; stream order and the tail text
# after the link (which strict readers were dropping) both survive; idempotent.
d = Document()
p = d.add_paragraph()
p._p.append(parse_xml(
    f'<w:r {nsdecls("w", "r")}><w:rPr><w:u w:val="single"/><w:color w:val="0645AD"/></w:rPr>'
    f'<w:t xml:space="preserve">See </w:t>'
    f'<w:hyperlink r:id="rId99" w:history="1"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr>'
    f'<w:t>the site</w:t></w:r></w:hyperlink>'
    f'<w:t xml:space="preserve"> for details.</w:t></w:r>'))
r1, out1 = run_unnest(d)
check(r1.paragraphs[0].text == "See the site for details.",
      "L1: stream changed: %r" % r1.paragraphs[0].text)
root1, links1 = _hyperlinks(out1)
check(len(links1) == 1, "L1: expected 1 hyperlink, got %d" % len(links1))
para1 = root1.find(f".//{W_MAIN}p")
check(any(child is links1[0] for child in para1),
      "L1: hyperlink is not a direct child of the paragraph")
out1b = de.hyperlink_unnest(out1)
check(out1b == out1, "L1: not idempotent")

# L2: directly adjacent same-rid fragments merge into one link; a different-rid
# neighbour stays separate.
d = Document()
p = d.add_paragraph()
for rid, txt in (("rId7", "speci"), ("rId7", "fication"), ("rId8", "elsewhere")):
    p._p.append(parse_xml(
        f'<w:r {nsdecls("w", "r")}><w:rPr><w:u w:val="single"/></w:rPr>'
        f'<w:hyperlink r:id="{rid}"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr>'
        f'<w:t>{txt}</w:t></w:r></w:hyperlink></w:r>'))
r2, out2 = run_unnest(d)
root2, links2 = _hyperlinks(out2)
check(len(links2) == 2, "L2: expected 2 hyperlinks after merge, got %d" % len(links2))
first_text = "".join(t.text or "" for t in links2[0].iter(f"{W_MAIN}t"))
check(first_text == "specification", "L2: merged link text %r" % first_text)
check(r2.paragraphs[0].text == "specificationelsewhere", "L2: stream changed")

# L3: rPr merge emits schema order even from two ordered inputs — wrapper
# [rFonts,color,u] + inner [rStyle,b,sz] must come out canonically ordered.
d = Document()
p = d.add_paragraph()
p._p.append(parse_xml(
    f'<w:r {nsdecls("w", "r")}><w:rPr><w:rFonts w:ascii="X"/><w:color w:val="FF0000"/>'
    f'<w:u w:val="single"/></w:rPr>'
    f'<w:hyperlink r:id="rId5"><w:r><w:rPr><w:rStyle w:val="Hyperlink"/>'
    f'<w:b/><w:sz w:val="28"/></w:rPr><w:t>x</w:t></w:r></w:hyperlink></w:r>'))
r3, out3 = run_unnest(d)
root3, links3 = _hyperlinks(out3)
inner_rpr = links3[0].find(f"{W_MAIN}r/{W_MAIN}rPr")
tags = [c.tag.split('}')[1] for c in inner_rpr]
check(tags == ["rStyle", "rFonts", "b", "color", "sz", "u"],
      "L3: merged rPr order %s" % tags)

# R23: a U+2010 compound broken at its own hyphen across a merge boundary is
# reconstructed exactly (no fuse, no injected space), while genuine
# hyphenation in the same document still heals. (iter-9; embedded font so the
# PDF text layer really carries U+2010.)
def _pdf_u2010(lines):
    pdf = fitz.open()
    pg = pdf.new_page(width=595, height=842)
    y = 80
    for t in lines:
        pg.insert_text((72, y), t, fontsize=11,
                       fontname="helv2", fontfile="/System/Library/Fonts/Helvetica.ttc")
        y += 15
    return pdf


CH1 = "The lab replaced all of its ageing equipment with modern state‐"
CH2 = "of‐the‐art spectrometers for the quarterly contamination analysis."
pdf = _pdf_u2010([CH1, CH2])
d = Document()
d.add_paragraph(CH1)
d.add_paragraph(CH2)
r, out = run_reflow(d, pdf)
body = " ".join(p.text for p in r.paragraphs if p.text)
check("state‐of‐the‐art" in body, "R23: compound chain not reconstructed: %r" % body[:90])
check("stateof" not in body.replace("‐", "").replace(" ", "")[:40] or "state‐of" in body,
      "R23: compound fused")

G1 = "A consistent grind and a level bed remove most of the equip‐"
G2 = "ment variability that beginners blame on the machine itself."
pdf = _pdf_u2010([G1, G2])
d = Document()
d.add_paragraph(G1)
d.add_paragraph(G2)
r, out = run_reflow(d, pdf)
body = " ".join(p.text for p in r.paragraphs if p.text)
check("equipment" in body, "R23b: genuine hyphenation no longer heals: %r" % body[:90])

print("hostile suite:", "ALL PASS" if not FAILS else f"{len(FAILS)} FAILURES")
sys.exit(1 if FAILS else 0)
