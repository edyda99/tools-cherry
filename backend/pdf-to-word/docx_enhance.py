"""Structural enhancement passes applied to pdf2docx output.

pdf2docx reproduces the look of a page but not Word's semantics: headings arrive as
big bold runs instead of Heading styles, list markers as frozen glyphs instead of
w:numPr, page furniture as body text instead of header/footer parts. Each pass here
upgrades one of those, working on the packed .docx bytes (plus the source PDF when a
pass needs per-page geometry). Passes are added one at a time by the quality loop,
each gated on the corpus scoreboard before it lands.

enhance() must stay safe on arbitrary documents: a pass that cannot prove its
transformation applies leaves the document unchanged, and a pass that raises is
skipped (logged as a metric line) rather than failing the conversion.
"""
import copy
import io
import json
import re
from collections import Counter

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

# A run is heading-sized when it clears the body size by 15% AND is either bold or
# dramatically larger. The bold requirement below 1.5x keeps emphasis-sized inline
# runs (a 14pt lead sentence in 11pt body) out of the outline.
HEAD_MIN_RATIO = 1.15
HEAD_BIG_RATIO = 1.5
HEAD_MAX_WORDS = 14
HEAD_MAX_LEVELS = 3
# Bail if the pass would style most of the document: that shape is a poster or
# cover page, not an outline, and spurious Heading styles are worse than none.
HEAD_MAX_SHARE = 0.6
HEAD_MAX_ABS = 40

_STYLE_XML = (
    '<w:style %s w:type="paragraph" w:styleId="Heading%d">'
    '<w:name w:val="heading %d"/><w:qFormat/>'
    '<w:pPr><w:keepNext/><w:keepLines/><w:outlineLvl w:val="%d"/></w:pPr>'
    "</w:style>"
)


def _run_size_pt(r):
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        return None
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        return None
    try:
        return float(sz.get(qn("w:val"))) / 2.0
    except (TypeError, ValueError):
        return None


def _run_bold(r):
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        return False
    b = rpr.find(qn("w:b"))
    return b is not None and (b.get(qn("w:val")) or "1").lower() not in ("0", "false", "none")


def _run_text(r):
    return "".join(t.text or "" for t in r.findall(qn("w:t")))


def _body_size_pt(paras):
    weight = {}
    for p in paras:
        for r in p.findall(qn("w:r")):
            t = _run_text(r).strip()
            sz = _run_size_pt(r)
            if t and sz:
                weight[sz] = weight.get(sz, 0) + len(t)
    if not weight:
        return None
    return max(weight.items(), key=lambda kv: kv[1])[0]


def _qualifies(sz, bold, body):
    if sz is None or sz < body * HEAD_MIN_RATIO:
        return False
    return bold or sz >= body * HEAD_BIG_RATIO


def _leading_heading_chunks(p, body):
    """(heading_chunks, rest_chunks, heading_size) over direct w:r / w:hyperlink
    children; whitespace-only runs ride along with whichever side they touch."""
    chunks = [c for c in p if c.tag in (qn("w:r"), qn("w:hyperlink"))]
    head, sizes = [], []
    for i, c in enumerate(chunks):
        if c.tag != qn("w:r") or _has_nontext_content(c):
            break  # hyperlinks, drawings, field chars: never part of a heading split
        t = _run_text(c)
        if not t.strip():
            head.append(c)
            continue
        sz = _run_size_pt(c)
        if _qualifies(sz, _run_bold(c), body):
            head.append(c)
            sizes.append(sz)
        else:
            break
    if not sizes:
        return [], chunks, None
    rest = [c for c in chunks if c not in head]
    return head, rest, max(sizes)


def _set_heading_style(p, level):
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        ppr = parse_xml("<w:pPr %s/>" % nsdecls("w"))
        p.insert(0, ppr)
    old = ppr.find(qn("w:pStyle"))
    if old is not None:
        ppr.remove(old)
    ppr.insert(0, parse_xml('<w:pStyle %s w:val="Heading%d"/>' % (nsdecls("w"), level)))


def _ensure_heading_styles(doc, levels):
    styles = doc.styles.element
    have = {s.get(qn("w:styleId")) for s in styles.findall(qn("w:style"))}
    for lvl in sorted(levels):
        if f"Heading{lvl}" not in have:
            styles.append(parse_xml(_STYLE_XML % (nsdecls("w"), lvl, lvl, lvl - 1)))


def heading_styles(data, pdf_doc=None):
    """Give heading-sized text real Heading styles, splitting headings that
    pdf2docx fused into the paragraph that follows them."""
    doc = Document(io.BytesIO(data))
    body_paras = [p._p for p in doc.paragraphs]
    body = _body_size_pt(body_paras)
    if not body:
        return data

    found = []  # (paragraph_element_to_style, heading_size)
    for p in body_paras:
        head, rest, hsize = _leading_heading_chunks(p, body)
        if not head:
            continue
        head_words = len(" ".join(_run_text(c) for c in head).split())
        if head_words == 0 or head_words > HEAD_MAX_WORDS:
            continue
        if not rest:
            found.append((p, hsize))
            continue
        rest_text = " ".join(_run_text(c) for c in rest if c.tag == qn("w:r"))
        if not rest_text.strip():
            found.append((p, hsize))
            continue
        # fused heading: pull the heading runs out into their own paragraph
        new_p = copy.deepcopy(p)
        for c in list(new_p):
            if c.tag in (qn("w:r"), qn("w:hyperlink")):
                new_p.remove(c)
        insert_at = len(new_p)  # after pPr, before nothing
        for c in head:
            p.remove(c)
            new_p.insert(insert_at, c)
            insert_at += 1
        p.addprevious(new_p)
        found.append((new_p, hsize))

    if not found:
        return data
    nonempty = sum(1 for p in body_paras if "".join(
        _run_text(r) for r in p.findall(qn("w:r"))).strip())
    # the share bail targets poster/cover shapes; tiny documents (a few
    # headings over little prose) are legitimate outlines, not posters
    if len(found) > HEAD_MAX_ABS or (nonempty >= 8 and len(found) / nonempty > HEAD_MAX_SHARE):
        return data

    distinct = sorted({round(sz * 2) / 2 for _, sz in found}, reverse=True)
    level_of = {sz: min(i + 1, HEAD_MAX_LEVELS) for i, sz in enumerate(distinct)}
    levels = set()
    for p, sz in found:
        lvl = level_of[round(sz * 2) / 2]
        _set_heading_style(p, lvl)
        levels.add(lvl)
    _ensure_heading_styles(doc, levels)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- list numbering ---------------------------------------------------------
# pdf2docx keeps list markers as literal text ("• ", "1. ") so Word sees plain
# paragraphs: no renumbering on edit, no continuation, no outline. This pass
# strips the frozen marker and attaches real w:numPr numbering.
#
# Safety model (each rule earned by a reproduced corruption in review):
# - Detection and stripping share ONE character stream: the text-like elements
#   of DIRECT w:r children. para.text is never used — it includes w:hyperlink
#   text the stripper cannot reach, which misaligns the two streams.
# - Only elements the stripper itself fully consumed are removed; a run is
#   removed only when the stripper emptied it AND nothing but rPr remains.
#   Runs carrying drawings, nested hyperlinks, field chars etc. are never
#   touched, and the python-docx run.text setter (which clears such children)
#   is never used.
# - Ordered stretches convert only when their printed numbers already read
#   1..n, every ordered level renders decimal "%N.", and one stretch gets one
#   shared ilvl — so Word can never show different numbers than the PDF did.
# - Bullet levels reuse the printed marker glyph as the level's lvlText.

BULLET_CHARS = "•●◦○▪·∙‣"
_BULLET_RE = re.compile(r"^\s*([" + BULLET_CHARS + r"\-–*])\s+")
_ORD_RE = re.compile(r"^\s*(\d{1,3})[.)]\s+")
LIST_LEVEL_GAP_PT = 12.0
LIST_MAX_LEVELS = 3
_BULLET_LVLTEXT = ("•", "◦", "▪")
_ORD_LVLS = (("decimal", "%1."), ("decimal", "%2."), ("decimal", "%3."))


def _indent_pt(para):
    ppr = para._p.find(qn("w:pPr"))
    if ppr is None:
        return 0.0
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        return 0.0
    try:
        return float(ind.get(qn("w:left")) or ind.get(qn("w:start"))) / 20.0
    except (TypeError, ValueError):
        return 0.0


_TEXTLIKE = (qn("w:t"), qn("w:tab"), qn("w:br"), qn("w:cr"), qn("w:noBreakHyphen"))


def _char_of(el):
    if el.tag == qn("w:t"):
        return el.text or ""
    if el.tag == qn("w:tab"):
        return "\t"
    if el.tag in (qn("w:br"), qn("w:cr")):
        return "\n"
    return "-"  # noBreakHyphen


def _has_nontext_content(r):
    return any(c.tag != qn("w:rPr") and c.tag not in _TEXTLIKE for c in r)


def _first_content_is_run(p):
    skip = {qn("w:pPr"), qn("w:proofErr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd"),
            qn("w:commentRangeStart"), qn("w:commentRangeEnd")}
    for c in p:
        if c.tag in skip:
            continue
        return c.tag == qn("w:r")
    return False


def _stream_text(p):
    """The exact character stream _consume_prefix can edit: text-like elements
    of DIRECT w:r children only (hyperlink/sdt content deliberately excluded)."""
    out = []
    for r in p.findall(qn("w:r")):
        for el in r:
            if el.tag in _TEXTLIKE:
                out.append(_char_of(el))
    return "".join(out)


def _consume_prefix(p, n):
    """Delete the first n stream characters. Touches only text-like elements,
    removes only elements it fully consumed, and removes a run only when it
    emptied it itself and nothing but rPr remains."""
    for r in list(p.findall(qn("w:r"))):
        if n <= 0:
            break
        touched = False
        for el in list(r):
            if n <= 0:
                break
            if el.tag == qn("w:t"):
                t = el.text or ""
                take = min(len(t), n)
                if not take:
                    continue
                n -= take
                touched = True
                rest = t[take:]
                if rest:
                    el.text = rest
                    el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                else:
                    r.remove(el)
            elif el.tag in _TEXTLIKE:
                n -= 1
                touched = True
                r.remove(el)
        if touched and not any(c.tag != qn("w:rPr") for c in r):
            p.remove(r)


def _numbering_root(doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    try:
        return doc.part.part_related_by(RT.NUMBERING).element
    except KeyError:
        return None


def _add_num(numbering, kind, bullet_chars=None):
    abs_ids = [int(a.get(qn("w:abstractNumId")) or 0)
               for a in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId")) or 0) for n in numbering.findall(qn("w:num"))]
    aid, nid = max(abs_ids, default=0) + 1, max(num_ids, default=0) + 1
    lvls = []
    for ilvl in range(LIST_MAX_LEVELS):
        if kind == "bul":
            fmt = "bullet"
            txt = (bullet_chars or {}).get(ilvl) or _BULLET_LVLTEXT[ilvl]
        else:
            fmt, txt = _ORD_LVLS[ilvl]
        lvls.append(f'<w:lvl w:ilvl="{ilvl}"><w:start w:val="1"/><w:numFmt w:val="{fmt}"/>'
                    f'<w:lvlText w:val="{txt}"/><w:lvlJc w:val="left"/></w:lvl>')
    abs_el = parse_xml(f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{aid}">'
                       f'<w:multiLevelType w:val="hybridMultilevel"/>{"".join(lvls)}</w:abstractNum>')
    num_el = parse_xml(f'<w:num {nsdecls("w")} w:numId="{nid}">'
                       f'<w:abstractNumId w:val="{aid}"/></w:num>')
    # CT_Numbering sequence: numPicBullet*, abstractNum*, num*, numIdMacAtCleanup?
    cleanup = numbering.find(qn("w:numIdMacAtCleanup"))
    nums = numbering.findall(qn("w:num"))
    abs_anchor = nums[0] if nums else cleanup
    if abs_anchor is not None:
        abs_anchor.addprevious(abs_el)
    else:
        numbering.append(abs_el)
    if cleanup is not None:
        cleanup.addprevious(num_el)
    else:
        numbering.append(num_el)
    return nid


def _set_numpr(para, ilvl, numid):
    # python-docx's CT_PPr accessors place numPr per the schema sequence
    ppr = para._p.get_or_add_pPr()
    numpr = ppr.get_or_add_numPr()
    numpr.get_or_add_ilvl().set(qn("w:val"), str(ilvl))
    numpr.get_or_add_numId().set(qn("w:val"), str(numid))


def _block_levels(indents):
    order = sorted({round(i, 1) for i in indents})
    levels, lvl, prev = {}, 0, None
    for ind in order:
        if prev is not None and ind - prev > LIST_LEVEL_GAP_PT:
            lvl += 1
        levels[ind] = min(lvl, LIST_MAX_LEVELS - 1)
        prev = ind
    return levels


def list_numbering(data, pdf_doc=None):
    """Turn frozen marker glyphs into real Word list numbering."""
    doc = Document(io.BytesIO(data))
    items = []  # (paragraph_index, para, kind, match, indent_pt)
    for i, para in enumerate(doc.paragraphs):
        if not _first_content_is_run(para._p):
            continue  # marker inside a hyperlink/sdt: not ours to edit
        text = _stream_text(para._p)
        m = _ORD_RE.match(text) or _BULLET_RE.match(text)
        if m is None or not text[m.end():].strip():
            continue  # no marker, or marker-only paragraph
        kind = "ord" if m.re is _ORD_RE else "bul"
        items.append((i, para, kind, m, _indent_pt(para)))
    if not items:
        return data

    # dash/asterisk "bullets" are ambiguous with prose; keep one only when an
    # adjacent paragraph is also a dash/asterisk bullet
    ambiguous = "-–*"
    by_idx = {it[0]: it for it in items}
    items = [it for it in items
             if not (it[2] == "bul" and it[3].group(1) in ambiguous)
             or any(n in by_idx and by_idx[n][2] == "bul"
                    and by_idx[n][3].group(1) in ambiguous
                    for n in (it[0] - 1, it[0] + 1))]
    if not items:
        return data

    blocks, cur = [], [items[0]]
    for it in items[1:]:
        if it[0] == cur[-1][0] + 1:
            cur.append(it)
        else:
            blocks.append(cur)
            cur = [it]
    blocks.append(cur)

    numbering = _numbering_root(doc)
    if numbering is None:
        return data

    convert = []  # (item, numid, ilvl)
    for block in blocks:
        levels = _block_levels([it[4] for it in block])
        bullet_nid = None
        bullets = [it for it in block if it[2] == "bul"]
        if bullets:
            # keep the printed glyph per level so the look doesn't change
            per_level = {}
            for it in bullets:
                per_level.setdefault(levels[round(it[4], 1)], []).append(it[3].group(1))
            glyphs = {lvl: Counter(chars).most_common(1)[0][0]
                      for lvl, chars in per_level.items()}
        i = 0
        while i < len(block):
            if block[i][2] == "ord":
                j = i
                while j < len(block) and block[j][2] == "ord":
                    j += 1
                seq = block[i:j]
                printed = [int(e[3].group(1)) for e in seq]
                # split at each printed "1": adjacent restarting lists (natural,
                # or made adjacent by furniture removal) convert per segment;
                # every segment must itself read 1..k or the stretch declines
                starts = [k for k, v in enumerate(printed) if v == 1]
                ok = bool(starts) and starts[0] == 0
                segments = []
                if ok:
                    bounds = starts + [len(seq)]
                    for a, b in zip(bounds, bounds[1:]):
                        if printed[a:b] != list(range(1, b - a + 1)):
                            ok = False
                            break
                        segments.append(seq[a:b])
                if ok:
                    for seg in segments:
                        # one segment = one level: per-item indent jitter must
                        # never split a printed 1..k sequence across counters
                        ilvl = Counter(levels[round(e[4], 1)]
                                       for e in seg).most_common(1)[0][0]
                        nid = _add_num(numbering, "ord")
                        convert.extend((e, nid, ilvl) for e in seg)
                i = j
            else:
                if bullet_nid is None:
                    bullet_nid = _add_num(numbering, "bul", glyphs)
                it = block[i]
                convert.append((it, bullet_nid, levels[round(it[4], 1)]))
                i += 1

    if not convert:
        return data
    for (idx, para, kind, m, ind), nid, ilvl in convert:
        _consume_prefix(para._p, m.end())
        _set_numpr(para, ilvl, nid)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- header / footer parts --------------------------------------------------
# pdf2docx writes page furniture (running headers, footers) into the body once
# per page, because a .docx has no page boundaries to hang it on. This pass uses
# the source PDF's geometry to move that furniture into real header/footer
# parts.
#
# Safety model (each rule earned by a reproduced corruption in review):
# - A text is furniture only when it sits in the top/bottom band on EVERY page
#   and matches EXACTLY page-count body paragraphs. Anything else declines:
#   cover pages, per-chapter headers, occurrences pdf2docx merged into a body
#   paragraph or a table would otherwise cause deleted content or half-removal.
# - A text qualifying in both bands declines (the header sweep would steal the
#   footer's occurrences).
# - Matched paragraphs must be plain text (pPr + text runs only): bookmarks,
#   fields, notes, links, drawings, numbering and section breaks never move
#   into a part or get deleted with one.
# - Documents already carrying header/footer machinery (a section owning a
#   header/footerReference, titlePg, evenAndOddHeaders) decline entirely —
#   which also makes the pass idempotent.

HF_BAND_FRAC = 0.12
HF_MIN_PAGES = 2


def _norm_furniture(s):
    return re.sub(r"\s+", " ", s or "").strip()


def _band_lines(pdf_doc):
    head, foot = {}, {}
    for page in pdf_doc:
        r = page.rect
        top = r.y0 + r.height * HF_BAND_FRAC
        bot = r.y1 - r.height * HF_BAND_FRAC
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                x0, y0, x1, y1 = line["bbox"]
                text = _norm_furniture("".join(s.get("text", "") for s in line.get("spans", [])))
                if not text:
                    continue
                if y1 <= top:
                    target = head
                elif y0 >= bot:
                    target = foot
                else:
                    continue
                e = target.setdefault(text, {"pages": set(), "y": y0})
                e["pages"].add(page.number)
                e["y"] = min(e["y"], y0)
    return head, foot


def _plain_furniture_para(p):
    """True when the paragraph is plain text: safe to delete and to clone into
    a header/footer part. Anything beyond pPr + text runs (bookmarks, fields,
    notes, links, drawings, numbering, section breaks) declines."""
    for c in p:
        if c.tag == qn("w:pPr"):
            if c.find(qn("w:sectPr")) is not None or c.find(qn("w:numPr")) is not None:
                return False
        elif c.tag == qn("w:r"):
            for rc in c:
                if rc.tag != qn("w:rPr") and rc.tag not in _TEXTLIKE:
                    return False
        else:
            return False
    return True


def _cell_texts(doc):
    out = set()
    for tc in doc.element.body.iter(qn("w:tc")):
        for p in tc.iter(qn("w:p")):
            text = _norm_furniture("".join(t.text or "" for t in p.iter(qn("w:t"))))
            if text:
                out.add(text)
    return out


def _hf_blocked(doc, kind):
    ref = qn(f"w:{kind}Reference")
    for sp in doc.element.body.iter(qn("w:sectPr")):
        if sp.find(ref) is not None or sp.find(qn("w:titlePg")) is not None:
            return True
    try:
        if doc.settings.element.find(qn("w:evenAndOddHeaders")) is not None:
            return True
    except Exception:
        pass
    return False


def header_footer_parts(data, pdf_doc=None):
    """Move repeated page furniture into real Word header/footer parts."""
    if pdf_doc is None or pdf_doc.page_count < HF_MIN_PAGES:
        return data
    pages = pdf_doc.page_count
    head, foot = _band_lines(pdf_doc)
    hdr_all = {t: i for t, i in head.items() if len(i["pages"]) == pages}
    ftr_all = {t: i for t, i in foot.items() if len(i["pages"]) == pages}
    overlap = set(hdr_all) & set(ftr_all)
    wanted = {
        "header": sorted((i["y"], t) for t, i in hdr_all.items() if t not in overlap),
        "footer": sorted((i["y"], t) for t, i in ftr_all.items() if t not in overlap),
    }
    if not wanted["header"] and not wanted["footer"]:
        return data

    doc = Document(io.BytesIO(data))
    cells = _cell_texts(doc)
    moved = {"header": [], "footer": []}
    for kind in ("header", "footer"):
        if not wanted[kind] or _hf_blocked(doc, kind):
            continue
        for y, text in wanted[kind]:
            matches = [p for p in doc.paragraphs if _norm_furniture(p.text) == text]
            if len(matches) != pages or text in cells:
                continue
            if not all(_plain_furniture_para(p._p) for p in matches):
                continue
            exemplar = copy.deepcopy(matches[0]._p)
            for p in matches:
                p._p.getparent().remove(p._p)
            moved[kind].append(exemplar)

    if not moved["header"] and not moved["footer"]:
        return data
    section = doc.sections[0]
    for kind, part in (("header", section.header), ("footer", section.footer)):
        if not moved[kind]:
            continue
        part.is_linked_to_previous = False  # fresh part: exactly one empty paragraph
        default_p = part.paragraphs[0]._p
        for exemplar in moved[kind]:
            default_p.addprevious(exemplar)
        if not "".join(t.text or "" for t in default_p.iter(qn("w:t"))).strip():
            default_p.getparent().remove(default_p)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- paragraph reflow --------------------------------------------------------
# pdf2docx fragments logical paragraphs wherever its block detection breaks:
# styled-run line boundaries, column jumps, page breaks. It also carries the
# PDF's end-of-line hyphenation into the text ("equip‐ment"). This pass merges
# fragments back and removes hyphenation — but ONLY where the source PDF
# testifies: two docx paragraphs merge only when their combined token stream is
# contiguous inside one PDF-derived logical paragraph, and a hyphen is removed
# only when the fused word appears at a PDF line break. Merging moves content,
# never deletes it; ambiguity is always a no-op.

HYPHENS = "-‐­"
_TERMINALS = ".!?:;。．！？"  # incl. fullwidth/CJK — pdf2docx splits blocks on these


def _tok(s):
    # unicode-aware: Cyrillic/Arabic/CJK text must be visible to the oracle,
    # or paragraphs carrying it would look empty and get merged over
    return re.findall(r"[^\W_]+", (s or "").lower())


def _pdf_logical_paras(pdf_doc):
    """Reading-order logical paragraphs + words fused across hyphenated line
    breaks. Lines inside a MuPDF block are segmented, not blindly joined: a
    line continues its paragraph only when the previous line is nearly full
    width, does not end a sentence into a capital, and the new line is not
    first-line indented. Segments then join across blocks/columns/pages under
    the same linguistic rules plus tight geometry. Every rule here exists
    because blind joining collapsed book layouts, poems, signature blocks and
    heading-only pages in adversarial review."""
    segs = []  # (page_no, bbox, text, font_size, page_y0, page_h, last_line_full)
    fuse_candidates = []  # (fused_word, hyphen_char)
    interior = set()  # hyphenated forms seen mid-line: NOT hyphenation artifacts

    def note_fuse(a_text, b_text):
        a_text = a_text.rstrip()
        w1 = re.findall(r"[A-Za-z]+", a_text[-24:])
        w2 = re.findall(r"[A-Za-z]+", b_text[:24])
        if not w1 or not w2:
            return
        # a hyphen inside the continuation word, or a chain of them behind the
        # boundary, marks a typographic compound broken at its own hyphen
        # (state-of-the-art) — never a hyphenation artifact (iter-9 guard)
        right_word = re.match(r"\S+", b_text.lstrip())
        left_word = re.search(r"\S+$", a_text)
        if right_word and any(h in right_word.group(0) for h in HYPHENS):
            return
        if left_word and sum(left_word.group(0).count(h) for h in "‐­") >= 2:
            return
        fuse_candidates.append(((w1[-1] + w2[0]).lower(), a_text[-1]))

    def continues(prev_text, prev_full, cur_text, cur_indented):
        a_last = prev_text.rstrip()[-1:]
        b_first = cur_text.lstrip()[:1]
        if not a_last or not b_first or not prev_full or cur_indented:
            return False
        if a_last in _TERMINALS:
            return False
        return b_first.islower() or (a_last in HYPHENS and b_first.isalpha())

    for page in pdf_doc:
        for b in page.get_text("dict").get("blocks", []):
            if b.get("type") != 0:
                continue
            lines = []  # (text, bbox, size)
            for ln in b.get("lines", []):
                t = "".join(s.get("text", "") for s in ln.get("spans", [])).strip()
                if t:
                    size = max((s.get("size", 10.0) for s in ln.get("spans", [])), default=10.0)
                    lines.append((t, ln["bbox"], size))
            if not lines:
                continue
            for (a, _, _), (nxt, _, _) in zip(lines, lines[1:]):
                if a[-1:] in tuple(HYPHENS) and nxt[:1].islower():
                    note_fuse(a, nxt)
            for t, _, _ in lines:
                for m in re.finditer(r"([A-Za-z]+)[" + HYPHENS + r"]([A-Za-z]+)", t):
                    interior.add((m.group(1) + m.group(2)).lower())
            bx0 = min(l[1][0] for l in lines)
            bx1 = max(l[1][2] for l in lines)
            bw = max(bx1 - bx0, 1.0)
            cur_lines, cur_size = [lines[0]], lines[0][2]
            for prev_l, cur_l in zip(lines, lines[1:]):
                # a continuation implies the previous line broke at the column
                # edge: it must fill its block AND the block must be a real
                # text column, not a stack of short standalone lines
                pw = prev_l[1][2] - prev_l[1][0]
                prev_full = pw >= 0.70 * bw and pw >= 90.0
                indented = cur_l[1][0] > bx0 + 0.5 * cur_l[2]
                if continues(prev_l[0], prev_full, cur_l[0], indented):
                    cur_lines.append(cur_l)
                else:
                    segs.append(_seg_of(page, cur_lines, bw, bx0))
                    cur_lines = [cur_l]
            segs.append(_seg_of(page, cur_lines, bw, bx0))

    # No joining ACROSS blocks: two adversarial reviews proved the geometry
    # gate cannot tell a paragraph break from a line break there (lowercase
    # unterminated paragraphs weld; page-boundary text glues to furniture).
    # A logical paragraph is exactly an intra-block segment, and only segments
    # from real text columns (>=140pt wide) may authorise a merge.
    paras = [seg[2] for seg in segs]
    eligible = [seg[8] for seg in segs]

    # Only typographic hyphens (U+2010 / soft hyphen) mark automatic line
    # breaks; an ASCII "-" at a line end is indistinguishable from a compound
    # broken on its own hyphen ("re-\nform" vs "reform") and never fuses. A
    # form the document also hyphenates mid-line is spelling, never an
    # artifact — meaning must not flip.
    fused = {w for w, h in fuse_candidates if h in "‐­" and w not in interior}

    logical = []
    for p in paras:
        toks = _tok(p)
        # apply the fuses the paragraph itself testified to
        out, k = [], 0
        while k < len(toks):
            if k + 1 < len(toks) and toks[k] + toks[k + 1] in fused:
                out.append(toks[k] + toks[k + 1])
                k += 2
            else:
                out.append(toks[k])
                k += 1
        logical.append(" ".join(out))
    return logical, fused, eligible


def _seg_of(page, lines, block_width, block_x0):
    x0 = min(l[1][0] for l in lines)
    y0 = min(l[1][1] for l in lines)
    x1 = max(l[1][2] for l in lines)
    y1 = max(l[1][3] for l in lines)
    lw = lines[-1][1][2] - lines[-1][1][0]
    last_full = lw >= 0.60 * block_width and lw >= 90.0
    starts_indented = lines[0][1][0] > block_x0 + 0.5 * lines[0][2]
    merge_ok = block_width >= 140.0
    return (page.number, (x0, y0, x1, y1), " ".join(l[0] for l in lines),
            lines[-1][2], page.rect.y0, page.rect.height, last_full, starts_indented,
            merge_ok)


def _reflow_mergeable(p):
    ppr = p.find(qn("w:pPr"))
    if ppr is None:
        return True
    return (ppr.find(qn("w:pStyle")) is None and ppr.find(qn("w:numPr")) is None
            and ppr.find(qn("w:sectPr")) is None)


_SPACER_PPR_OK = None  # built lazily: qn() needs the docx namespace map loaded


def _is_blank_spacer(p):
    """A spacer may be swept with a merge ONLY when it carries nothing at all:
    no visible text in any script, no drawings, breaks, bookmarks or fields —
    and no pPr machinery beyond cosmetic spacing (a pStyle, pageBreakBefore or
    framePr changes layout and must survive)."""
    global _SPACER_PPR_OK
    if _SPACER_PPR_OK is None:
        _SPACER_PPR_OK = {qn("w:rPr"), qn("w:spacing"), qn("w:jc"), qn("w:ind")}
    for c in p:
        if c.tag == qn("w:pPr"):
            if any(g.tag not in _SPACER_PPR_OK for g in c):
                return False
            continue
        if c.tag != qn("w:r"):
            return False
        for rc in c:
            if rc.tag == qn("w:rPr"):
                continue
            # ascii-whitespace strip only: an NBSP is content, not blankness
            if rc.tag != qn("w:t") or (rc.text or "").strip(" \t\r\n"):
                return False
    return True


def _strip_trailing_hyphen(p):
    ts = [t for t in p.iter(qn("w:t")) if t.text and t.text.rstrip()]
    if ts:
        txt = ts[-1].text.rstrip()
        if txt[-1:] in tuple(HYPHENS):
            ts[-1].text = txt[:-1]


def _dehyph_within(doc, fused):
    # no optional space: the interior guard sees no spaced forms, so the fuser
    # must not rewrite them either
    pat = re.compile(r"([A-Za-z]{2,})[" + HYPHENS + r"]([A-Za-z]{2,})")
    changed = False
    for para in doc.paragraphs:
        if not _reflow_mergeable(para._p):
            continue  # headings, list items, section paragraphs keep their spelling
        runs = para._p.findall(qn("w:r"))
        for r in runs:
            for t in r.findall(qn("w:t")):
                if t.text and any(h in t.text for h in HYPHENS):
                    new = pat.sub(
                        lambda m: m.group(1) + m.group(2)
                        if (m.group(1) + m.group(2)).lower() in fused else m.group(0),
                        t.text)
                    if new != t.text:
                        t.text = new
                        changed = True
        # hyphen at a run boundary: "…equip‐" + "ment…"
        for ra, rb in zip(runs, runs[1:]):
            ta = [t for t in ra.findall(qn("w:t")) if t.text and t.text.rstrip()]
            tb = [t for t in rb.findall(qn("w:t")) if t.text and t.text.strip()]
            if not ta or not tb:
                continue
            atxt, btxt = ta[-1].text.rstrip(), tb[0].text.lstrip()
            if atxt[-1:] in tuple(HYPHENS):
                w1 = re.findall(r"[A-Za-z]+", atxt[-24:])
                w2 = re.findall(r"[A-Za-z]+", btxt[:24])
                if w1 and w2 and (w1[-1] + w2[0]).lower() in fused:
                    ta[-1].text = atxt[:-1]
                    tb[0].text = tb[0].text.lstrip()
                    changed = True
    return changed


def paragraph_reflow(data, pdf_doc=None):
    """Merge paragraph fragments back together, guided by the source PDF."""
    if pdf_doc is None:
        return data
    logical, fused, eligible = _pdf_logical_paras(pdf_doc)
    if not logical:
        return data
    wrapped = [" " + lp + " " for lp in logical]
    full = set(logical)

    doc = Document(io.BytesIO(data))
    changed = False
    cursor = 0  # logical paragraphs are consumed in document order:
    i = 0       # a fragment pair may never match an earlier region's text
    while True:
        paras = doc.paragraphs
        if i >= len(paras):
            break
        A = paras[i]
        la = _tok(A.text)
        if not la:
            i += 1
            continue
        # next content paragraph; only paragraphs carrying NOTHING (no text in
        # any script, no drawings/breaks/bookmarks) count as sweepable spacers
        spacers, B, j = [], None, i + 1
        while j < len(paras):
            if _is_blank_spacer(paras[j]._p):
                spacers.append(paras[j])
                j += 1
                continue
            B = paras[j]
            break
        if B is None:
            break
        sa = " ".join(la)
        if sa in full:
            for k in range(cursor, len(logical)):
                if logical[k] == sa:
                    cursor = k + 1
                    break
            i = j
            continue
        chain = [A._p] + [s._p for s in spacers] + [B._p]
        if (not all(chain[k].getnext() is chain[k + 1] for k in range(len(chain) - 1))
                or not (_reflow_mergeable(A._p) and _reflow_mergeable(B._p))):
            i = j
            continue
        lb = _tok(B.text)
        if not lb:
            i = j
            continue
        variants = [" " + " ".join(la + lb) + " "]
        a_end = A.text.rstrip()[-1:]
        dehyph = a_end in HYPHENS and (la[-1] + lb[0]) in fused
        if dehyph:
            variants.insert(0, " " + " ".join(la[:-1] + [la[-1] + lb[0]] + lb[1:]) + " ")
        hit = next((v for v in variants
                    if any(v in wrapped[k] for k in range(cursor, len(logical))
                           if eligible[k])), None)
        if hit is None:
            i = j
            continue
        if dehyph and hit is variants[0]:
            _strip_trailing_hyphen(A._p)
        else:
            stream_a = _stream_text(A._p)
            # a word-attached hyphen at the seam joins directly: the compound
            # keeps its printed form (state-of-the-art, well-known) instead of
            # gaining a space after the hyphen (iter-9)
            hyphen_join = bool(re.search(r"[A-Za-z][-‐­]$", stream_a)) and B.text[:1].isalpha()
            if (not hyphen_join and not stream_a.endswith((" ", "\t", "\n"))
                    and not B.text[:1].isspace()):
                joiner = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve"> </w:t></w:r>')
                last_run = A._p.findall(qn("w:r"))
                if last_run:
                    rpr = last_run[-1].find(qn("w:rPr"))
                    if rpr is not None:
                        rpr = copy.deepcopy(rpr)
                        for deco in ("w:u", "w:strike", "w:dstrike", "w:shd",
                                     "w:highlight", "w:em", "w:bdr", "w:vertAlign"):
                            el = rpr.find(qn(deco))
                            if el is not None:
                                rpr.remove(el)
                        joiner.insert(0, rpr)
                A._p.append(joiner)
        for c in list(B._p):
            if c.tag != qn("w:pPr"):
                A._p.append(c)
        B._p.getparent().remove(B._p)
        for s in spacers:
            s._p.getparent().remove(s._p)
        changed = True
        # stay on A: it may continue absorbing the next fragment

    changed = _dehyph_within(doc, fused) or changed
    if not changed:
        return data
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- span-boundary space repair ----------------------------------------------
# pdf2docx drops the inter-word space where a styled or hyperlink span meets
# plain text ("with thedeployment checklist"). A space is restored at a run
# boundary only under double evidence from the PDF's own word stream: the
# fused form is NOT a word the PDF contains, AND the two halves DO appear
# adjacent as separate words. Insertion-only; ambiguity is a no-op.

_EDGE_PUNCT = re.compile(r"^[^\w]+|[^\w]+$")
# ASCII letters/digits only: the seam string tested MUST be the seam string
# edited. Punctuation at a seam ("on"+"-call", "O'"+"Brien", "ISO"+"-9001")
# always declines — stripping it first made the guard test a different string
# than the document contains. CJK seams decline too (no inter-word spaces).
_TAIL_WORD = re.compile(r"[A-Za-z0-9]+$")
_HEAD_WORD = re.compile(r"^[A-Za-z0-9]+")


def _pdf_word_evidence(pdf_doc):
    words, bigrams = set(), set()
    for page in pdf_doc:
        seq = [w for w in (_EDGE_PUNCT.sub("", t[4]).lower()
                           for t in page.get_text("words", sort=True)) if w]
        words.update(seq)
        bigrams.update(zip(seq, seq[1:]))
    return words, bigrams


def _seam_stream(p):
    """Text-like elements of a paragraph in document order, without descending
    into drawings/text boxes/fallback content — their text is not body text."""
    skip = {qn("w:drawing"), qn("w:object"), qn("w:pict")}
    out = []

    def walk(el):
        for c in el:
            if c.tag in skip:
                continue
            if c.tag in _TEXTLIKE:
                out.append(c)
            else:
                walk(c)

    walk(p)
    return out


def span_space_repair(data, pdf_doc=None):
    """Restore inter-word spaces pdf2docx loses at span boundaries."""
    if pdf_doc is None:
        return data
    words, bigrams = _pdf_word_evidence(pdf_doc)
    if not bigrams:
        return data
    doc = Document(io.BytesIO(data))
    changed = False
    for para in doc.paragraphs:
        stream = _seam_stream(para._p)
        for a, b in zip(stream, stream[1:]):
            if a.tag != qn("w:t") or b.tag != qn("w:t"):
                continue
            ta, tb = a.text or "", b.text or ""
            if not ta or not tb:
                continue
            if not (ta[-1].isalnum() and tb[0].isalnum()):
                continue  # only letter-against-letter seams can be lost spaces
            m1, m2 = _TAIL_WORD.search(ta), _HEAD_WORD.search(tb)
            if not m1 or not m2:
                continue
            f1, f2 = m1.group().lower(), m2.group().lower()
            # the WHOLE seam token must be the tested fragment: an interior
            # hyphen or non-ASCII letter ("X-Ray"+"scanner", "Zürich"+"bank")
            # would make the veto probe a substring of a token the PDF holds
            # solid, and corrupt it
            full_tail = re.search(r"\S+$", ta).group()
            full_head = re.search(r"^\S+", tb).group()
            if (_EDGE_PUNCT.sub("", full_tail).lower() != f1
                    or _EDGE_PUNCT.sub("", full_head).lower() != f2):
                continue
            if (f1 + f2) not in words and (f1, f2) in bigrams:
                a.text = ta + " "
                a.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                changed = True
    if not changed:
        return data
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- wrap_break_heal: remove pdf2docx's per-wrapped-line hard breaks ----------
# pdf2docx writes a w:br for every wrapped source line it does not space-join,
# so the converted paragraph never re-wraps when edited, and end-of-line
# hyphenations stay frozen mid-word. Each break is healed only when the source
# PDF testifies the boundary is a line WRAP: the previous line fills its block
# (>=70% of block width AND >=90pt) and the block is wide (>=140pt). Poems,
# addresses and other intentional short lines fail the fullness gate and are
# kept. Edits are element-local: a br becomes a single space (or is deleted
# when whitespace already surrounds it), and a typographic line-break hyphen
# (U+2010/U+00AD) fuses its halves unless the same hyphenated form also occurs
# mid-line in the source (then it is a real compound and the break is kept).

_WRAP_FULL_SHARE = 0.70
_WRAP_FULL_MIN_PT = 90.0
_WRAP_BLOCK_MIN_PT = 140.0
_SOFT_HYPHENS = "‐­"
_SKIP_SUBTREES = (qn("w:drawing"), qn("w:object"), qn("w:pict"), qn("w:hyperlink"),
                  "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent")


def _wb_blocks(pdf_doc):
    blocks = []
    for page in pdf_doc:
        for b in page.get_text("dict")["blocks"]:
            if b.get("type") != 0:
                continue
            lines = []
            for ln in b.get("lines", []):
                t = "".join(s["text"] for s in ln["spans"])
                if t.strip():
                    lines.append({"text": t, "x0": ln["bbox"][0], "x1": ln["bbox"][2]})
            if lines:
                blocks.append({"lines": lines,
                               "x0": min(l["x0"] for l in lines),
                               "x1": max(l["x1"] for l in lines)})
    return blocks


def _wb_interior_forms(blocks):
    """(left, right) pairs of hyphenated forms seen MID-line: real compounds
    the fusion path must never join."""
    forms = set()
    for b in blocks:
        for ln in b["lines"]:
            text = ln["text"]
            for m in re.finditer(r"([A-Za-z]+)[-‐­]([A-Za-z]+)", text):
                if m.end() < len(text.rstrip()):
                    forms.add((m.group(1).lower(), m.group(2).lower()))
    return forms


_CJK_RANGES = ((0x3000, 0x30FF), (0x3400, 0x4DBF), (0x4E00, 0x9FFF),
               (0xF900, 0xFAFF), (0xFF00, 0xFFEF), (0xAC00, 0xD7AF))


def _wb_cjk(ch):
    return any(a <= ord(ch) <= b for a, b in _CJK_RANGES)


def _wb_items(p):
    """Text-like elements of the paragraph in stream order, with their parent
    run — skipping drawing/object/pict/hyperlink/AlternateContent subtrees so
    nothing inside them is ever counted or edited."""
    items = []

    def walk(el):
        for c in el:
            if c.tag in _SKIP_SUBTREES or c.tag == qn("w:pPr"):
                continue
            if c.tag in _TEXTLIKE:
                items.append(c)
            else:
                walk(c)

    walk(p)
    return items


def _wb_segments(items):
    """Split the item stream at wrap w:br elements. Returns (segments, brs):
    segments is a list of lists of w:t/w:tab/... elements, brs the separating
    br elements (len(segments) == len(brs) + 1). None if the paragraph has any
    br the pass must not touch (page/column breaks)."""
    segments, brs, cur = [], [], []
    for el in items:
        if el.tag == qn("w:br"):
            if el.get(qn("w:type")) not in (None, "textWrapping"):
                return None, None
            brs.append(el)
            segments.append(cur)
            cur = []
        elif el.tag == qn("w:cr"):
            return None, None
        else:
            cur.append(el)
    segments.append(cur)
    return segments, brs


def _wb_text(seg):
    return "".join(_char_of(el) for el in seg)


def _wb_tokens(s):
    return s.split()


def _wb_align(segments, block):
    """Greedily map each segment to one or more consecutive block lines by
    exact token equality. Returns the index of the LAST line of each segment,
    or None when the paragraph does not line up with this block."""
    ends, li = [], 0
    lines = block["lines"]
    for seg in segments:
        want = _wb_tokens(_wb_text(seg))
        if not want:
            return None
        got, start = [], li
        while li < len(lines) and len(got) < len(want):
            got.extend(_wb_tokens(lines[li]["text"]))
            li += 1
        if got != want or li == start:
            return None
        ends.append(li - 1)
    if li != len(lines):
        return None
    return ends


def _wb_drop_empty(el):
    run = el.getparent()
    run.remove(el)
    if not [k for k in run if k.tag != qn("w:rPr")] and not _has_nontext_content(run):
        run.getparent().remove(run)


def _wb_strip_trailing(seg, chars):
    """Remove trailing whitespace plus one char of `chars` from the segment's
    element stream (fusion must leave no stray space: 'pro ‐' + 'cedure' has to
    become 'procedure'). Whitespace-only w:t elements passed on the way are
    deleted; self-emptied elements and runs are removed."""
    trailing_ws = []
    for el in reversed(seg):
        if el.tag != qn("w:t"):
            if el.tag == qn("w:tab"):
                return False
            continue
        txt = el.text or ""
        if not txt.strip():
            trailing_ws.append(el)
            continue
        stripped = txt.rstrip()
        if stripped[-1] not in chars:
            return False
        el.text = stripped[:-1].rstrip()
        if not el.text:
            _wb_drop_empty(el)
        for ws in trailing_ws:
            _wb_drop_empty(ws)
        return True
    return False


def _wb_strip_leading_ws(seg):
    for el in seg:
        if el.tag != qn("w:t"):
            return
        txt = el.text or ""
        if txt.strip():
            el.text = txt.lstrip()
            return
        _wb_drop_empty(el)


def _wb_remove_br(br, replace_with_space):
    run = br.getparent()
    if replace_with_space:
        sp = parse_xml('<w:t xml:space="preserve" %s> </w:t>' % nsdecls("w"))
        run.replace(br, sp)
        return
    run.remove(br)
    if not [k for k in run if k.tag != qn("w:rPr")] and not _has_nontext_content(run):
        run.getparent().remove(run)


def wrap_break_heal(data, pdf_doc=None):
    if pdf_doc is None:
        return data
    blocks = _wb_blocks(pdf_doc)
    if not blocks:
        return data
    interior = _wb_interior_forms(blocks)
    doc = Document(io.BytesIO(data))
    body = doc.element.body
    changed = False

    for p in body.findall(qn("w:p")):
        items = _wb_items(p)
        if not any(el.tag == qn("w:br") for el in items):
            continue
        segments, brs = _wb_segments(items)
        if not brs:
            continue
        seg_texts = [_wb_text(s) for s in segments]
        match = None
        for block in blocks:
            ends = _wb_align(segments, block)
            if ends is not None:
                match = (block, ends)
                break
        if match is None:
            continue
        block, ends = match
        if (block["x1"] - block["x0"]) < _WRAP_BLOCK_MIN_PT:
            continue

        for i, br in enumerate(brs):
            prev = block["lines"][ends[i]]
            full = ((prev["x1"] - block["x0"]) >= _WRAP_FULL_SHARE * (block["x1"] - block["x0"])
                    and (prev["x1"] - prev["x0"]) >= _WRAP_FULL_MIN_PT)
            if not full:
                continue
            left = seg_texts[i].rstrip()
            right = seg_texts[i + 1].lstrip()
            if not left or not right:
                continue
            if (seg_texts[i].rstrip(" ").endswith("\t")
                    or seg_texts[i + 1].lstrip(" ").startswith("\t")):
                continue  # a tab at the seam is structure, not a wrap
            if left[-1] in _SOFT_HYPHENS:
                m1 = re.search(r"([A-Za-z]+)[‐­]$", left)
                m2 = re.match(r"([A-Za-z]+)", right)
                if not m1 or not m2:
                    continue
                if (m1.group(1).lower(), m2.group(1).lower()) in interior:
                    continue
                right_word = re.match(r"\S+", right).group(0)
                left_word = re.search(r"\S+$", left).group(0)
                # a hyphen inside the continuation word, or a chain of them on
                # the left, marks a real typographic compound (state-of-the-art)
                if (any(h in right_word for h in "‐­-")
                        or left_word.count("‐") + left_word.count("­") >= 2):
                    continue
                if not _wb_strip_trailing(segments[i], _SOFT_HYPHENS):
                    continue
                _wb_strip_leading_ws(segments[i + 1])
                _wb_remove_br(br, replace_with_space=False)
                changed = True
            elif left[-1] == "-":
                _wb_remove_br(br, replace_with_space=False)
                changed = True
            else:
                needs_space = (not seg_texts[i][-1:].isspace()
                               and not seg_texts[i + 1][:1].isspace()
                               and not _wb_cjk(left[-1]) and not _wb_cjk(right[0]))
                _wb_remove_br(br, replace_with_space=needs_space)
                changed = True

    if not changed:
        return data
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_HYPERLINK_STYLE_XML = (
    '<w:style %s w:type="character" w:styleId="Hyperlink">'
    '<w:name w:val="Hyperlink"/><w:basedOn w:val="DefaultParagraphFont"/>'
    '<w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr></w:style>'
)


def _link_key(h):
    return (h.get(qn("r:id")), h.get(qn("w:anchor")))


_RPR_ORDER = ("rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps", "strike",
              "dstrike", "outline", "shadow", "emboss", "imprint", "noProof", "snapToGrid",
              "vanish", "webHidden", "color", "spacing", "w", "kern", "position", "sz",
              "szCs", "highlight", "u", "effect", "bdr", "shd", "fitText", "vertAlign",
              "rtl", "cs", "em", "lang", "eastAsianLayout", "specVanish", "oMath")
_RPR_IDX = {qn("w:%s" % t): i for i, t in enumerate(_RPR_ORDER)}


def _rpr_insert(rpr, el):
    my = _RPR_IDX.get(el.tag, len(_RPR_IDX))
    for child in rpr:
        if _RPR_IDX.get(child.tag, len(_RPR_IDX)) > my:
            child.addprevious(el)
            return
    rpr.append(el)


def _merge_wrapper_rpr(wrapper_rpr, inner_run):
    """Give an inner run the wrapper run's direct formatting. pdf2docx puts the
    visible underline/color on the wrapper and only rStyle on the inner run, so
    on the (never observed) tag conflict the WRAPPER's value wins — it is what
    Word was rendering. Inner-only properties are inserted at their CT_RPr
    schema position."""
    if wrapper_rpr is None:
        return
    merged = copy.deepcopy(wrapper_rpr)
    old = inner_run.find(qn("w:rPr"))
    if old is not None:
        for child in old:
            if merged.find(child.tag) is None:
                _rpr_insert(merged, copy.deepcopy(child))
        inner_run.remove(old)
    inner_run.insert(0, merged)


def hyperlink_unnest(data, pdf_doc=None):
    """Lift w:hyperlink elements that pdf2docx nests INSIDE w:r up to their
    schema-valid position as siblings of the run. Word tolerates the invalid
    nesting, but schema-strict consumers (LibreOffice, QuickLook, and other
    non-Word apps) drop the whole subtree, deleting the link text on screen.

    Only moves nodes within their own container in document order — the
    character stream is unchanged. Also merges directly-adjacent fragments of
    the same link and defines the referenced Hyperlink character style, which
    pdf2docx names but never defines."""
    doc = Document(io.BytesIO(data))
    changed = False

    work = [h for h in doc.element.body.xpath(".//w:hyperlink")
            if h.getparent().tag == qn("w:r")]
    while work:
        link = work.pop(0)
        run = link.getparent()
        if run is None or run.tag != qn("w:r"):
            continue
        changed = True
        container = run.getparent()
        wrapper_rpr = run.find(qn("w:rPr"))

        kids = list(run)
        at = kids.index(link)
        tail = [k for k in kids[at + 1:]]

        run.remove(link)
        container.insert(list(container).index(run) + 1, link)
        for inner in link.findall(qn("w:r")):
            _merge_wrapper_rpr(wrapper_rpr, inner)
        if tail:
            tail_run = parse_xml("<w:r %s/>" % nsdecls("w"))
            if wrapper_rpr is not None:
                tail_run.append(copy.deepcopy(wrapper_rpr))
            for k in tail:
                run.remove(k)
                tail_run.append(k)
            container.insert(list(container).index(link) + 1, tail_run)
            work[:0] = [h for h in tail_run.iter(qn("w:hyperlink"))
                        if h.getparent().tag == qn("w:r")]
        if not [k for k in run if k.tag != qn("w:rPr")]:
            container.remove(run)
        if container.tag == qn("w:r"):
            work.append(link)

    if changed:
        for link in doc.element.body.xpath(".//w:hyperlink"):
            prev = link.getprevious()
            while (prev is not None and prev.tag == qn("w:hyperlink")
                   and _link_key(prev) == _link_key(link) and _link_key(link) != (None, None)):
                for k in list(link):
                    prev.append(k)
                parent = link.getparent()
                parent.remove(link)
                link = prev
                prev = link.getprevious()

    has_links = bool(doc.element.body.xpath(".//w:hyperlink"))
    if has_links:
        styles_el = doc.styles.element
        defined = any(s.get(qn("w:styleId")) == "Hyperlink"
                      for s in styles_el.findall(qn("w:style")))
        if not defined:
            styles_el.append(parse_xml(_HYPERLINK_STYLE_XML % nsdecls("w")))
            changed = True

    if not changed:
        return data
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Order is load-bearing and enhance() is single-shot: hyperlink_unnest runs
# first so every later pass sees schema-valid hyperlink positions, and
# span_space_repair must see the document BEFORE reflow's dehyphenation (a
# healed word looks like a lost-space seam to a second run). The pipeline
# calls enhance() exactly once per conversion; never chain it.
# wrap_break_heal is NOT enabled: adversarial review (iter 8) proved geometric
# evidence cannot separate line wraps from deliberate lines (addresses,
# signatures, code, logs, TOCs, verse-only pages) — see backlog #8 dead end 1
# and out/adv_iter8b. The machinery stays for a combined-evidence attempt #2.
PASSES = (hyperlink_unnest, span_space_repair, header_footer_parts, heading_styles,
          list_numbering, paragraph_reflow)


def enhance(docx_bytes, pdf_doc=None):
    """All accepted passes, in order. A failing pass is skipped, never fatal."""
    data = docx_bytes
    for pass_fn in PASSES:
        try:
            data = pass_fn(data, pdf_doc)
        except Exception as e:  # noqa: BLE001 - conversion must survive a bad pass
            print(json.dumps({"m": "enhance_pass_failed", "pass": pass_fn.__name__,
                              "err": f"{type(e).__name__}: {e}"[:300]}))
    return data
